from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.apps import apps
from io import BytesIO


def get_course_model():
    """Localiza o modelo de disciplina no app courses."""
    for name in ('Disciplina', 'Discipline'):
        try:
            return apps.get_model('courses', name)
        except LookupError:
            continue
    return None


def _get_enrollments(course):
    """Tenta localizar a relação de matrículas usando nomes comuns."""
    for attr in (
        'enrollments',
        'enrollment_set',
        'students',
        'registrations',
        'subscriptions',
    ):
        if hasattr(course, attr):
            qs = getattr(course, attr)
            try:
                return qs.select_related('student__user')
            except Exception:
                try:
                    return qs.all()
                except Exception:
                    return qs
    return []


def _extract_user_from_enrollment(e):
    """Retorna um objeto user-like a partir de um objeto enrollment,
    tentando vários padrões."""
    if (student := getattr(e, 'student', None)) is not None:
        if (user := getattr(student, 'user', None)):
            return user
        if hasattr(student, 'get_full_name') or hasattr(student, 'email'):
            return student

    if (user := getattr(e, 'user', None)):
        return user

    full_name = getattr(e, 'full_name', '') or getattr(e, 'name', '')
    username = getattr(e, 'username', '')
    email = getattr(e, 'email', '')
    if full_name or username or email:
        class _SimpleUser:
            def __init__(self, full_name, username, email):
                self._full_name = full_name
                self.username = username
                self.email = email

            def get_full_name(self):
                return self._full_name

        return _SimpleUser(full_name, username, email)

    return None


def _enrollment_number(e):
    for name in (
        'enrollment_number', 'registration_number', 'matricula', 'id'
    ):
        if (val := getattr(e, name, None)):
            return val
    return ''


def _pdf_escape_text(value):
    text = (
        str(value or '')
        .replace('\\', '\\\\')
        .replace('(', '\\(')
        .replace(')', '\\)')
    )
    return text.encode('cp1252', errors='replace').decode('latin-1')


def _build_simple_pdf(lines, title):
    page_width = 842
    page_height = 595
    left_margin = 36
    top_margin = 32
    line_height = 12
    max_lines_per_page = 42

    pages = [lines[i:i + max_lines_per_page] for i in range(0, len(lines), max_lines_per_page)]
    if not pages:
        pages = [[]]

    objects = []

    def add_object(data):
        objects.append(data)
        return len(objects)

    font_obj = add_object(b"<< /Type /Font /Subtype /Type1 /BaseFont /Courier >>")

    page_ids = []
    content_ids = []
    for page_lines in pages:
        content_lines = [
            "BT",
            "/F1 10 Tf",
            f"1 0 0 1 {left_margin} {page_height - top_margin} Tm",
            f"({_pdf_escape_text(title)}) Tj",
        ]
        y_offset = line_height * 2
        for line in page_lines:
            content_lines.append(
                f"1 0 0 1 {left_margin} {page_height - top_margin - y_offset} Tm"
            )
            content_lines.append(f"({_pdf_escape_text(line)}) Tj")
            y_offset += line_height
        content_lines.append("ET")
        content_stream = "\n".join(content_lines).encode("latin-1")
        content_obj = add_object(
            b"<< /Length " + str(len(content_stream)).encode() + b" >>\nstream\n"
            + content_stream + b"\nendstream"
        )
        content_ids.append(content_obj)
        page_ids.append(add_object(b""))

    pages_obj = add_object(b"")
    for index, page_obj in enumerate(page_ids):
        page_dict = (
            f"<< /Type /Page /Parent {pages_obj} 0 R /MediaBox [0 0 {page_width} {page_height}] "
            f"/Resources << /Font << /F1 {font_obj} 0 R >> >> /Contents {content_ids[index]} 0 R >>"
        ).encode()
        objects[page_obj - 1] = page_dict

    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects[pages_obj - 1] = (
        f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>"
    ).encode()

    catalog_obj = add_object(f"<< /Type /Catalog /Pages {pages_obj} 0 R >>".encode())

    buffer = BytesIO()
    buffer.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for object_id, obj in enumerate(objects, start=1):
        offsets.append(buffer.tell())
        buffer.write(f"{object_id} 0 obj\n".encode())
        buffer.write(obj)
        buffer.write(b"\nendobj\n")

    xref_offset = buffer.tell()
    buffer.write(f"xref\n0 {len(objects) + 1}\n".encode())
    buffer.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        buffer.write(f"{offset:010d} 00000 n \n".encode())

    buffer.write(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_obj} 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF"
        ).encode()
    )
    return buffer.getvalue()


def export_enrollment_doc(request, course_id):
    """Gera um .docx com a lista de alunos matriculados."""
    try:
        from docx import Document
    except ImportError:
        return HttpResponse(
            'Dependência faltando: python-docx. Execute: '
            'pip install python-docx',
            status=500,
        )

    CourseModel = get_course_model()
    if not CourseModel:
        return HttpResponse('Modelo de disciplina não encontrado.', status=500)

    course = get_object_or_404(CourseModel, pk=course_id)
    enrollments = _get_enrollments(course)

    heading = 'Lista de alunos - '
    if hasattr(course, 'nome'):
        heading += f'{course.nome}'
    if hasattr(course, 'area'):
        heading += f' - {course.area}'
    doc = Document()
    doc.add_heading(heading, level=1)

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nome'
    hdr_cells[1].text = 'Nome de usuário'
    hdr_cells[2].text = 'E-mail'
    hdr_cells[3].text = 'Matrícula'

    for e in enrollments:
        user = _extract_user_from_enrollment(e)
        row_cells = table.add_row().cells
        row_cells[0].text = user.get_full_name() if user else ''
        row_cells[1].text = getattr(user, 'username', '') if user else ''
        row_cells[2].text = getattr(user, 'email', '') if user else ''
        row_cells[3].text = str(_enrollment_number(e))

    filename = 'matriculados'
    if hasattr(course, 'nome'):
        filename += f'_{course.nome}'
    response = HttpResponse(
        content_type=(
            'application/vnd.openxmlformats-officedocument.'
            'wordprocessingml.document'
        )
    )
    response['Content-Disposition'] = f'attachment; filename={filename}.docx'
    doc.save(response)
    return response


def export_enrollment_pdf(request, course_id):
    """Gera um PDF com a lista de alunos matriculados."""

    CourseModel = get_course_model()
    if not CourseModel:
        return HttpResponse('Modelo de disciplina não encontrado.', status=500)

    course = get_object_or_404(CourseModel, pk=course_id)
    enrollments = _get_enrollments(course)

    try:
        from weasyprint import HTML

        html_string = render_to_string(
            'students/reports/enrollment_pdf.html',
            {'course': course, 'enrollments': enrollments},
        )
        html = HTML(string=html_string, base_url=request.build_absolute_uri('/'))
        pdf = html.write_pdf()
    except ImportError:
        course_name = getattr(course, 'nome', 'Disciplina')
        area_name = getattr(getattr(course, 'area', None), 'nome', '') or getattr(course, 'area', '')
        title = f'Lista de alunos - {course_name}'
        if area_name:
            title = f'{title} - {area_name}'

        lines = [
            'Nome | Nome de usuário | E-mail | Matrícula',
            '-' * 110,
        ]
        for enrollment in enrollments:
            user = _extract_user_from_enrollment(enrollment)
            full_name = user.get_full_name() if user else ''
            username = getattr(user, 'username', '') if user else ''
            email = getattr(user, 'email', '') if user else ''
            matricula = _enrollment_number(enrollment)
            lines.append(
                f'{full_name[:30]:30} | {username[:20]:20} | {email[:35]:35} | {str(matricula)[:12]}'
            )

        pdf = _build_simple_pdf(lines, title=title)

    filename = 'matriculados'
    if hasattr(course, 'nome'):
        filename += f'_{course.nome}'
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename={filename}.pdf'
    return response
